from flask import Flask, render_template, request, session


from flask import Flask, render_template, request, session, flash, redirect, url_for 
import os

from flask_wtf import FlaskForm
from wtforms import StringField, FloatField, SubmitField
from wtforms.validators import DataRequired
from flask_pymongo import PyMongo
from pymongo import MongoClient
from docx import Document

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your_secret_key'
app.config['MONGO_URI'] = "mongodb+srv://ekkoisademon:sOiVkni05wTw0I9b@cluster0.a9nzbwg.mongodb.net/"
mongo = PyMongo(app)

def connect_to_mongodb():
    client = MongoClient(app.config['MONGO_URI'])
    db = client['database_IPD']
    return db

def import_data(db):
    data = {
        "material_inventory": list(db.material_inventory.find()),
        "recipes": list(db.recipes.find()),
        "human_labor_surcharges": list(db.human_labor_surcharges.find()),
        "machine_inventory": list(db.machine_inventory.find()),
        "billing_orders": list(db.billing_orders.find())
    }
    print("Data imported successfully.")
    return data

def create_bill_from_form(items, db, data):
    print("Creating bill from form...")

    # Initialize an empty list to store bill items
    bill_items = []

    for item in items:
        recipe_name = item["recipe_name"].lower()
        amount = float(item["amount"])  # Convert amount to float

        # Retrieve the recipe from the data
        recipe = next((r for r in data["recipes"] if r["recipe_name"].lower() == recipe_name), None)

        if recipe:
            # Check if 'recipe_cost' is available and is a valid number
            if "recipe_cost" in recipe and isinstance(recipe["recipe_cost"], (int, float)):
                # Calculate the cost for the given amount
                item_cost = recipe["recipe_cost"] * amount
                total_cost = item_cost * 1.12  # 12% tax

                bill_item = {"recipe_name": recipe_name, "amount": amount, "item_cost": item_cost}
                bill_items.append(bill_item)

                # Uncomment the following lines to print the bill item details
                # print(f"Bill item details: {bill_item}")
            else:
                print(f"Invalid or missing 'recipe_cost' for recipe: {recipe_name}")

    if bill_items:
        total_cost = sum(item["item_cost"] for item in bill_items) * 1.12
        bill = {"items": bill_items, "total_cost": total_cost}

        # Update session with the new bill items
        session['bill_items'] = bill_items

        db.billing_orders.insert_one(bill)
        destock_based_on_bill(bill, data, db)  # Pass db to destock_based_on_bill
        doc_file = export_bill_as_doc(bill)
        send_bill_to_mongodb(bill, doc_file, db)

        print("Bill generated successfully.")
    else:
        print("No valid items found in the form data.")


def destock_based_on_bill(bill, data, db):
    print("Destocking material based on the bill...")
    for item in bill["items"]:
        recipe = next((r for r in data["recipes"] if r["recipe_name"].lower() == item["recipe_name"]), None)
        if recipe:
            for material, ratio in recipe["materials_required"].items():
                used_amount = ratio * item["amount"]
                material_entry = next((m for m in data["material_inventory"] if m["name"] == material), None)
                if material_entry:
                    material_entry["current_amount_kg"] -= used_amount

    # Ensure that the material_inventory collection is updated in MongoDB
    db.material_inventory.drop()
    db.material_inventory.insert_many(data["material_inventory"])
    print("Material inventory updated successfully.")

# Import the bill_counter module
from bill_counter import get_next_bill_number

def export_bill_as_doc(bill):
    # Get the next bill number
    bill_number = get_next_bill_number()
    # Specify the file path for the bills folder
    bills_folder = 'bills/'
    # Create the bills folder if it doesn't exist
    os.makedirs(bills_folder, exist_ok=True)
    # Construct the file name with the bill number
    file_name = f'billing_order_{bill_number}.docx'
    # Construct the full file path
    file_path = os.path.join(bills_folder, file_name)
    document = Document()
    document.add_heading('Billing Order', 0)

    document.add_heading('Items:', level=1)
    for item in bill["items"]:
        document.add_paragraph(f'{item["recipe_name"]}: {item["amount"]} kg')

    document.add_heading('Total Cost:', level=2)
    document.add_paragraph(f'${bill["total_cost"]:.2f}')
    # Save the document to the bills folder
    document.save(file_path)

    print(f"Bill exported as document successfully. Bill Number: {bill_number}")
    return file_path


def send_bill_to_mongodb(bill, doc_file, db):
    with open(doc_file, 'rb') as file_content:
        document_data = {"file_content": file_content.read()}
        # Ensure that the billing_orders collection is updated in MongoDB
        db.billing_orders.update_one({"_id": bill["_id"]}, {"$set": {"document": document_data}})
    print("Bill data sent to MongoDB successfully.")


# Updated restock_material_from_form function
def restock_material_from_form(extracted_data, data, db):
    # Iterate through the extracted data
    print(extracted_data)
    for i, entry in enumerate(extracted_data['material_names']):
        material_name = entry
        amount = float(extracted_data['amounts'][i]) if extracted_data['amounts'][i] else 0.0

        # Update the material entry directly in MongoDB
        db.material_inventory.update_one(
            {'name': material_name},
            {'$inc': {'current_amount_kg': amount}}
        )

    print("Material inventory updated successfully.")


# GUI button order
buttons_order = [
    ('Home', 'homepage'),
    ('Material Inventory', 'material_inventory'),
    ('Recipes', 'recipes'),
    ('Human Labor Surcharges', 'human_labor_surcharges'),
    ('Machine Inventory', 'machine_inventory'),
    ('Billing Orders', 'billing_orders'),
    ('Generate Bill', 'generate_bill'),
    ('Restock', 'restock_material'),
    ('Reorder', 'display_items_below_reorder')
]

@app.route('/')
def homepage():
    db = connect_to_mongodb()
    data = import_data(db)
    return render_template('index.html', data=data, buttons_order=buttons_order)

@app.route('/material_inventory')
def material_inventory():
    db = connect_to_mongodb()
    data = import_data(db)
    return render_template('material_inventory.html', data=data, buttons_order=buttons_order)

@app.route('/recipes')
def recipes():
    db = connect_to_mongodb()
    data = import_data(db)
    return render_template('recipes.html', data=data, buttons_order=buttons_order)

@app.route('/human_labor_surcharges')
def human_labor_surcharges():
    db = connect_to_mongodb()
    data = import_data(db)
    return render_template('human_labor_surcharges.html', data=data, buttons_order=buttons_order)

@app.route('/machine_inventory')
def machine_inventory():
    db = connect_to_mongodb()
    data = import_data(db)
    return render_template('machine_inventory.html', data=data, buttons_order=buttons_order)

@app.route('/billing_orders')
def billing_orders():
    db = connect_to_mongodb()
    data = import_data(db)
    return render_template('billing_orders.html', data=data, buttons_order=buttons_order)

# Example form class (modify based on your actual form class)
class GenerateBillForm(FlaskForm):
    recipe_name = StringField('Recipe Name')
    amount = FloatField('Amount (kg)')
    submit = SubmitField('Generate Bill')

# Route to handle bill generation
@app.route('/generate_bill', methods=['GET', 'POST'])
def generate_bill():
    form = GenerateBillForm()

    # Print the raw form data
    print("Raw Form Data:", request.form)

    db = connect_to_mongodb()
    data = import_data(db)
    recipe_table_data = [(recipe['recipe_name'], recipe['recipe_cost']) for recipe in data['recipes']]

    # Extract non-empty recipe names and their corresponding amounts
    form_data = request.form.to_dict(flat=False)
    extracted_data = {
        'recipe_names': [name.lower() for name in form_data.get('recipe_name[]', []) if name.strip()],
        'amounts': form_data.get('amount[]', [])
    }

    # Clear the session bill_items before adding new items
    session['bill_items'] = []

    # Iterate through the extracted data and add items to the session
    for i, recipe_name in enumerate(extracted_data['recipe_names']):
        amount = extracted_data['amounts'][i]

        print(f"Processing Item {i + 1} - Recipe Name: {recipe_name}, Amount: {amount}")

        # No validation, add the item directly to the session
        bill_item = {"recipe_name": recipe_name, "amount": amount}
        session['bill_items'].append(bill_item)

    print("Bill items in session:", session['bill_items'])

    if session['bill_items']:
        print("Creating bill from form...")
        create_bill_from_form(session['bill_items'], db, data)
        session['bill_items'] = []  # Reset session after creating the bill
        print("Bill created successfully.")
    else:
        print("No valid items found in the form data.")

    # Enumerate the bill items for displaying in the template
    bill_items_enum = list(enumerate(session.get('bill_items', [])))

    # Set the maximum number of items
    max_items = 25

    print("Rendering template...")
    return render_template('generate_bill.html', form=form, recipe_table_data=recipe_table_data,
                           bill_items=session.get('bill_items', []), data=data,
                           max_items=max_items,  # Include max_items here
                           bill_items_enum=bill_items_enum)



class RestockMaterialForm(FlaskForm):
    material_name = StringField('Material Name', validators=[DataRequired()])
    amount = FloatField('Amount (kg)', validators=[DataRequired()])
    submit = SubmitField('Restock Material')


# Updated app route for /restock_material
@app.route('/restock_material', methods=['GET', 'POST'])
def restock_material():
    form = RestockMaterialForm()
    db = connect_to_mongodb()
    data = import_data(db)

    # Assuming you have a function to get material inventory data from MongoDB
    material_inventory_data = [(material['name'], material['current_amount_kg']) for material in db.material_inventory.find()]

    if request.method == 'POST':
        try:
            # Extracted data
            extracted_data = {
                'material_names': [name for name in request.form.getlist('material_name[]') if name.strip()],
                'amounts': [float(amount) for amount in request.form.getlist('amount[]')]
            }
  


            print("Extracted Data:", extracted_data)

            # Update in-memory data
            restock_material_from_form(extracted_data, data, db)

            # Import updated data after restocking
            data = import_data(db)

            flash('Material restocked successfully!', 'success')
            return redirect(url_for('restock_material'))

        except Exception as e:
            print(f'Error processing restock request: {e}')
            flash(f'Error processing restock request: {e}', 'error')

    print("Rendering template...")

    # Render the form template with updated material inventory data
    return render_template('restock_material.html', form=form, material_inventory_data=material_inventory_data, data=data)


if __name__ == '__main__':
    app.run(debug=True)
